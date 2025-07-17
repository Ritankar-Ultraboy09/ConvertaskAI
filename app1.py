from flask import Flask, render_template, request, jsonify, url_for, redirect, send_file, json
from flask_sqlalchemy import SQLAlchemy
from waitress import serve
import os
from datetime import datetime
from before_whisper import convert_to_wav
from after_whisper import clean_text
from whisper_deploy import transcribe_chunk
from utils.chunking import chunk_audio_with_overlap
import shutil
import tempfile
from docx import Document
import io
from pydub import AudioSegment
import math
from TalkConvertask import Convertask_AI
import re
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
import time
from io import BytesIO
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from bs4 import BeautifulSoup
import markdown2


from summarisation import process_transcript_for_summary
from notion_push import push_to_notion

from email_initiate import EmailAgent

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = "postgresql+psycopg2://newaifyuser:user2@db:5432/new_db"
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

email_agent = EmailAgent()







class AudioFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(10000), nullable=False)
    mimetype = db.Column(db.String(500))
    transcript = db.Column(db.Text)
    upload_time = db.Column(db.DateTime, server_default=db.func.now())


class ModelFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(10000), nullable=False)
    Modeloutput = db.Column(db.Text)
    processed_time = db.Column(db.DateTime, server_default=db.func.now())

class Ai_output(db.Model):
    __tablename__ = "transcript_logs"
    id = db.Column(db.Integer, primary_key = True)
    transcript = db.Column(db.Text, nullable = False)
    ai_response = db.Column(db.Text, nullable = False)
    timestamp = db.Column(db.DateTime, default=db.func.now())

class EmailLog(db.Model):
    __tablename__ = "email_logs"
    id = db.Column(db.Integer, primary_key=True)
    subject = db.Column(db.String(500), nullable=False)
    recipients = db.Column(db.Text, nullable=False) 
    body = db.Column(db.Text, nullable=False)
    sent_successfully = db.Column(db.Boolean, default=False)
    timestamp = db.Column(db.DateTime, default=db.func.now())
    related_filename = db.Column(db.String(1000)) 





    



@app.route('/')
def index():
    return render_template("Convertask.html")

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No selected file'}), 400

    save_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(save_path)

    audio_file = AudioFile(filename=file.filename, mimetype=file.mimetype)
    db.session.add(audio_file)
    db.session.commit()


    return jsonify({
        'success': True,
        'filename': file.filename,
    })

@app.route('/transcribe', methods=['POST'])
def transcribe_file():
    data = request.get_json()
    filename = data.get('filename')
    if not filename:
        return jsonify({'success': False, 'error': 'Filename not provided'}), 400

    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    base_name = os.path.splitext(filename)[0]
    wav_path = os.path.join("uploads", f"{base_name}_converted.wav")
    
    try:
        start = time.time()
        print("[INFO] Starting conversion to WAV")
        convert_to_wav(input_path, wav_path)

        print("[INFO] Starting chunking")
        chunks = chunk_audio_with_overlap(wav_path)

        transcripts = []
        for i, chunk in enumerate(chunks):
            print(f"[INFO] Transcribing chunk {i + 1}/{len(chunks)}: {chunk}")
            transcripts.append(transcribe_chunk(chunk))
            os.remove(chunk)

        print("[INFO] Cleaning final transcript")
        full_text = clean_text(" ".join(transcripts))

        audio_record = AudioFile.query.filter_by(filename=filename).first()
        if audio_record:
            audio_record.transcript = full_text
            db.session.commit()

        print(f"Time Taken = {time.time()-start}")
        


        print("[INFO] Transcription complete")
        return jsonify({'success': True,'redirect_url': url_for('display_transcript', transcript=full_text, original_filename = filename)})

    except Exception as e:
        print(f"[ERROR] {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    


 
    
@app.route('/display_transcript')
def display_transcript():
    transcript = request.args.get('transcript', '') 
    original_filename = request.args.get('original_filename', '')
    return render_template('t_display.html', transcript=transcript, original_filename=original_filename)

@app.route('/update_transcript', methods=['POST'])
def update_transcript():
    data = request.get_json()
    edited_transcript = data.get('transcript')
    original_filename = data.get('original_filename') # Get the original filename

    if not edited_transcript or not original_filename:
        return jsonify({'success': False, 'error': 'Missing transcript or filename'}), 400

    try:
        audio_record = AudioFile.query.filter_by(filename=original_filename).first()
        if audio_record:
            audio_record.transcript = edited_transcript
            db.session.commit()
            return jsonify({'success': True, 'message': 'Transcript updated successfully'})
        else:
            return jsonify({'success': False, 'error': 'Audio file not found'}), 404
    except Exception as e:
        print(f"[ERROR] Failed to update transcript: {str(e)}")
        db.session.rollback() # Rollback in case of error
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500
    
@app.route('/download_transcript', methods=['POST'])
def download_transcript():
    
    data = request.get_json()
    transcript_text = data.get('transcript', '')
    download_format = data.get('format', '')
    filename = data.get('filename')

    if not transcript_text:
        return jsonify({'success': False, 'error': 'No transcript provided.'}), 400

    if download_format == 'docx':
        
        document = Document()
        document.add_paragraph(transcript_text)

        
        byte_stream = io.BytesIO()
        document.save(byte_stream)
        byte_stream.seek(0) 

        return send_file(
            byte_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}_transcript.docx'
        )
    elif download_format == 'pdf':
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph(transcript_text, styles['Normal']))
        doc.build(story)
        buffer.seek(0)

        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{filename}_transcript.pdf'
        )
    else:
        return jsonify({'success': False, 'error': 'Invalid download format requested.'}), 400
    

@app.route('/summarise_transcript', methods=['POST'])
def summarise_transcript():
    print("Button is hitting")
    data = request.get_json()
    transcript_text = data.get('transcript', '')
    user_prompt = data.get("user_prompt", "").strip()
    original_filename = data.get('original_filename', 'unknown_audio')


    if not transcript_text:
        return jsonify({'success': False, 'error': 'No transcript provided.'}), 400
    
    try:
        end = time.time()
        print("[INFO] Calling structured summarization module...")
        summary_results = process_transcript_for_summary(transcript_text, user_prompt=user_prompt)
        print("[INFO] Structured summarization complete.")

        if "error" in summary_results:
            return jsonify({'success': False, 'error': summary_results["error"]}), 500
        
        print("Pushing to Notion")
        notion_result = push_to_notion(summary_results, speaker="Summary Model")
        print("Notion Push Complete")

        
        summary_json_string = json.dumps(summary_results, indent=2)

        new_model_file_entry = ModelFile(
            filename=original_filename,
            Modeloutput=summary_json_string
        )

        db.session.add(new_model_file_entry)
        db.session.commit()
        print(f"[INFO] Summarization results for {original_filename} committed to ModelFile database.")
        print(f"Time taken:- {time.time()-end}")

        return jsonify({
            'success': True,
            'redirect_url': url_for(
                'display_summary',
                summary_data=summary_json_string, 
                original_filename=original_filename 
            )
        })
    except Exception as e:
        print(f"[ERROR] Fatal error during summarization route: {str(e)}")
        import traceback
        traceback.print_exc() 
        return jsonify({'success': False, 'error': f"Server error during summarization: {str(e)}"}), 500
    
@app.route('/Talk_to_Convertask', methods= ['POST'])
def TalkToConvertask():
    print("Button is hitting")
    data = request.get_json()
    transcript_text = data.get('transcript', '')
    user_prompt_AI = data.get("user_prompt_1", "").strip()

    send_email = data.get('send_email', False)
    email_recipients = data.get('email_recipients', [])

    print(f"{user_prompt_AI}")

    if not transcript_text:
        return jsonify({'success': False, 'error': 'No transcript provided.'}), 400

    try:
        print("Connecting to ConvertaskAI")
        result = Convertask_AI(transcript_text, user_prompt_1=user_prompt_AI)
        plan_output = result.get("plan") or result.get("response", "")


        if "error" in result:
            return jsonify({'success': False, 'error': result["error"]}), 500
        
        
        return jsonify({
            'success': True, 
            'redirect_url': url_for('show_chat_response', output_text = plan_output)
        })

    except Exception as e:
        print(f"[ERROR] Fatal error during summarization route: {str(e)}")
        import traceback
        traceback.print_exc() 
        return jsonify({'success': False, 'error': f"Server error during summarization: {str(e)}"}), 500
    
def extract_email_from_message(message):
    
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, message)
    return emails[0] if emails else None

def detect_email_request(message):
    
    email_keywords = [
        'send me this plan on my email',
        'send this to my email',
        'email me this',
        'send me via email',
        'email this to me',
        'send this plan to email',
        'email me the plan',
        'send to email',
        'send me this to my email',
        'send this thing to my email',
        'email me',
        'send me this on my email'
    ]
    
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in email_keywords)

def create_formatted_pdf(content):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import black, darkblue
    import re
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch, bottomMargin=1*inch)
    styles = getSampleStyleSheet()
    

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        textColor=darkblue,
        alignment=0  
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=15,
        textColor=darkblue
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=10,
        textColor=darkblue
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        leftIndent=20,
        spaceAfter=6,
        bulletIndent=10
    )
    
    story = []
    
    
    story.append(Paragraph("Convertask AI ", title_style))
    story.append(Spacer(1, 12))
    
    lines = content.split('\n')
    current_list_items = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        
        if line.startswith('### '):
            
            if current_list_items:
                for item in current_list_items:
                    story.append(Paragraph(f"• {item}", bullet_style))
                current_list_items = []
            
            story.append(Paragraph(line[4:], subheading_style))
            
        elif line.startswith('## '):
            if current_list_items:
                for item in current_list_items:
                    story.append(Paragraph(f"• {item}", bullet_style))
                current_list_items = []
                
            story.append(Paragraph(line[3:], heading_style))
            
        elif line.startswith('# '):
            if current_list_items:
                for item in current_list_items:
                    story.append(Paragraph(f"• {item}", bullet_style))
                current_list_items = []
                
            story.append(Paragraph(line[2:], title_style))
            
        elif line.startswith('- ') or line.startswith('* '):
            current_list_items.append(line[2:])
            
        elif re.match(r'^\d+\.\s', line):
            if current_list_items:
                for item in current_list_items:
                    story.append(Paragraph(f"• {item}", bullet_style))
                current_list_items = []
            
            # Handle numbered lists
            story.append(Paragraph(line, styles['Normal']))
            
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            if current_list_items:
                for item in current_list_items:
                    story.append(Paragraph(f"• {item}", bullet_style))
                current_list_items = []
            
            story.append(Paragraph(f"<b>{line[2:-2]}</b>", styles['Normal']))
            
        else:
            if current_list_items:
                for item in current_list_items:
                    story.append(Paragraph(f"• {item}", bullet_style))
                current_list_items = []
                
            if line:
                story.append(Paragraph(line, styles['Normal']))
    
    if current_list_items:
        for item in current_list_items:
            story.append(Paragraph(f"• {item}", bullet_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer
def create_formatted_docx(content):
    """Create a well-formatted DOCX from markdown content"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    import re
    
    document = Document()
    
    
    title = document.add_heading('Convertask AI Strategic Plan', 0)
    title.alignment = 0  # Left alignment
    
    
    sentences = re.split(r'(?<=[.!?])\s+', content)
    
    
    section_headers = [
        'Vision / Long-term Goals',
        'Key Projects or Focus Areas', 
        'Immediate Action Plan or Roadmap',
        'Resource Requirements or Team Roles',
        'Funding Strategy or Investor Readiness',
        'Smart Next Steps or Questions'
    ]
    
    current_section = None
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        
        for header in section_headers:
            if header in sentence:
                
                header_start = sentence.find(header)
                if header_start != -1:
                    document.add_heading(header, level=1)
                    current_section = header
                    
                    remaining = sentence[header_start + len(header):].strip()
                    if remaining and remaining not in [':', '.', ',']:
                        document.add_paragraph(remaining)
                    break
        else:
           
            if sentence:
                
                if re.match(r'^[-*•]\s', sentence) or sentence.startswith('- '):
                    
                    clean_text = re.sub(r'^[-*•]\s*', '', sentence)
                    document.add_paragraph(clean_text, style='List Bullet')
                elif re.match(r'^\d+\.\s', sentence):
                    
                    document.add_paragraph(sentence, style='List Number')
                elif ':' in sentence and len(sentence.split(':')) == 2:
                    
                    parts = sentence.split(':', 1)
                    p = document.add_paragraph()
                    run = p.add_run(parts[0] + ':')
                    run.bold = True
                    p.add_run(' ' + parts[1])
                else:
                    
                    document.add_paragraph(sentence)
    
    
    if not any(header in content for header in section_headers):
        
        document = Document()
        document.add_heading('Convertask AI Strategic Plan', 0)
        
        
        paragraphs = re.split(r'(?:[.!?])\s+(?=[A-Z])', content)
        
        for para in paragraphs:
            para = para.strip()
            if para and len(para) > 10:  
                document.add_paragraph(para)
                document.add_paragraph()  
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

    
@app.route('/chat_continue', methods=['POST'])
def chat_continue():
    data = request.get_json()
    message = data.get('message', '')
    context = data.get('context', [])
    
    if not message:
        return jsonify({'success': False, 'error': 'No message provided'}), 400
    
    try:
        # Check if user is requesting email functionality
        if detect_email_request(message):
            email_address = extract_email_from_message(message)
            
            if not email_address:
                return jsonify({
                    'success': True,
                    'response': "I'd be happy to send you this plan via email! However, I need your email address. Please provide your email address in the format: 'Send me this plan on my email: your.email@example.com'"
                })
            
            
            ai_content = "\n\n".join(
                msg['content'] for msg in context if msg['role'] == 'assistant'
            )
            
            if not ai_content:
                return jsonify({
                    'success': True,
                    'response': "I don't have any recent plan or content to send. Please generate a plan first, then request to send it via email."
                })
            

            try:
                subject = "Your Convertask AI Plan"
                email_body = f"""
Hello,

Here's your strategic plan generated by Convertask AI:

{ai_content}

Best regards,

"""
                pdf_buffer = create_formatted_pdf(ai_content)
                docx_buffer = create_formatted_docx(ai_content)


               

                attachments = [
                    ("Convertask_Plan.pdf", pdf_buffer.read()),
                    ("Convertask_Plan.docx", docx_buffer.read())
                ]
                
                
                success = email_agent.send_email(
                    to_emails=[email_address],
                    subject=subject,
                    body=email_body,
                    attachments=attachments
                )
                
                # Log the email attempt
                email_log = EmailLog(
                    subject=subject,
                    recipients=email_address,
                    body=email_body,
                    sent_successfully=success,
                    related_filename="chat_plan"
                )
                db.session.add(email_log)
                db.session.commit()
                
                if success:
                    return jsonify({
                        'success': True,
                        'response': f"✅ Great! I've successfully sent your plan to {email_address}. Please check your inbox (and spam folder just in case)."
                    })
                else:
                    return jsonify({
                        'success': True,
                        'response': f"❌ Sorry, I encountered an issue while sending the email to {email_address}. Please check the email address and try again, or contact support if the issue persists."
                    })
                    
            except Exception as email_error:
                print(f"[ERROR] Email sending failed: {str(email_error)}")
                return jsonify({
                    'success': True,
                    'response': f"❌ Sorry, I encountered a technical issue while sending the email to {email_address}. Please try again later or contact support."
                })
        
       
        conversation_context = ""
        for msg in context[-5:]:  
            role = "User" if msg['role'] == 'user' else "Assistant"
            conversation_context += f"{role}: {msg['content']}\n\n"
        
        continue_prompt = f"""
You are an excellent AI Agent, who works and solves all the problems on the basis of 

Previous conversation context:
{conversation_context}

User's new message: {message}

Provide a helpful, strategic response that:
- Builds on the previous conversation context
- Addresses the user's specific question or request
- Offers actionable insights and recommendations
- Maintains a professional but conversational tone
- Uses clear, well-structured markdown format

⚠️ CRITICAL INSTRUCTIONS:
- NO explanations, NO thinking process, NO reasoning
- NO phrases like "First, I need to", "Let me break this down", "Putting this together"
- Start IMMEDIATELY with the structured output
- Be direct and actionable only

Focus on being practical, strategic, and directly helpful. Avoid unnecessary explanations of your process.
"""
        
        from TalkConvertask import Convertask
        response = Convertask(continue_prompt)
        
        if response is None:
            return jsonify({'success': False, 'error': 'Failed to get AI response'}), 500
        
       
        new_log = Ai_output(
            transcript=message,
            ai_response=response
        )
        db.session.add(new_log)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'response': response
        })
        
    except Exception as e:
        print(f"[ERROR] Chat continue error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Server error: {str(e)}'}), 500
    


################################################################################################### CHAT RESPONSE #####################################################################################################################




@app.route('/chat_response')
def show_chat_response():
    output_text = request.args.get('output_text', '')
    return render_template("AI_DISPLAY.html", output=output_text)

@app.route('/display_summary')
def display_summary():
    summary_data_json_string = request.args.get('summary_data', '{}')
    original_filename = request.args.get('original_filename', 'Unknown')
    try:
        summary_data = json.loads(summary_data_json_string)
    except json.JSONDecodeError:
        summary_data = {"error": "Invalid summary data format. Could not parse JSON."}
        print("[ERROR] Failed to decode summary_data_json_string:", summary_data_json_string)

    return render_template('summary.html', summary_data=summary_data, original_filename=original_filename)






@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))



if __name__ == "__main__":
    print("Runs only when app1.py")
    with app.app_context():
        serve(app, host='0.0.0.0', port=8000)


    
    